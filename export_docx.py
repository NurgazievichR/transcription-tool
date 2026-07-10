"""Build client-style transcript .docx (speaker / timestamp / source / empty translation)."""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


HEADERS = (
    "Speaker Name",
    "Timestamp",
    "Source Language Transcript",
    "English Translation",
)


def format_docx_timestamp(seconds: float) -> str:
    """Match reference style: 0:00:08.7"""
    if seconds < 0:
        seconds = 0.0
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    return f"{h}:{m:02d}:{s:04.1f}"


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, *, bold: bool = False, size_pt: float = 11, color: RGBColor | None = None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color


def build_transcript_docx(utterances: list[dict], *, title: str | None = None) -> bytes:
    """
    Create a 4-column table like the client reference docx.

    Columns: Speaker | Timestamp | Source | English Translation (empty).
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    if title:
        heading = doc.add_paragraph()
        run = heading.add_run(title)
        _set_run_font(run, bold=True, size_pt=14)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = True

    header_row = table.rows[0]
    for i, label in enumerate(HEADERS):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(label)
        _set_run_font(run, bold=True, size_pt=11, color=RGBColor(255, 255, 255))
        _set_cell_shading(cell, "2F5496")

    for utt in utterances:
        row = table.add_row()
        speaker = utt.get("speaker_label") or utt.get("speaker") or "N/A"
        ts = format_docx_timestamp(float(utt.get("start", 0)))
        source = (utt.get("text") or "").strip()
        values = (speaker, ts, source, "")

        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            _set_run_font(run, bold=(i == 0), size_pt=10.5)

    # Prefer readable column widths
    widths = (Cm(3.2), Cm(2.8), Cm(8.5), Cm(5.5))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
